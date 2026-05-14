"""Document text extractor for AI coach attachments.

Extracts text from .docx, .xlsx, .txt, .md files for LLM consumption.
"""
from __future__ import annotations

import logging
from io import BytesIO

_log = logging.getLogger(__name__)

DOCUMENT_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.document",
    "text/plain",
    "text/markdown",
}

MAX_TEXT_LENGTH = 15000


def extract_document_text(raw: bytes, mime_type: str, filename: str) -> str:
    """Extract text content from a document. Returns structured text description."""
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(raw, filename)
    elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.document":
        return _extract_xlsx(raw, filename)
    elif mime_type in ("text/plain", "text/markdown"):
        return _extract_text(raw, filename)
    raise ValueError(f"Unsupported document MIME type: {mime_type}")


def _truncate(text: str) -> str:
    if len(text) > MAX_TEXT_LENGTH:
        return text[:MAX_TEXT_LENGTH] + "\n...[内容过长已截断]"
    return text


def _extract_docx(raw: bytes, filename: str) -> str:
    from docx import Document

    doc = Document(BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text: list[str] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables_text.append("\n".join(rows))

    parts = [f"[Word文档: {filename}]"]
    if paragraphs:
        parts.append("\n".join(paragraphs))
    if tables_text:
        parts.append("[表格内容]\n" + "\n\n".join(tables_text))
    return _truncate("\n\n".join(parts))


def _extract_xlsx(raw: bytes, filename: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    try:
        parts = [f"[Excel表格: {filename}]"]
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(max_row=200, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
        return _truncate("\n\n".join(parts))
    finally:
        wb.close()


def _extract_text(raw: bytes, filename: str) -> str:
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("gbk", errors="replace")
    ext_label = "Markdown文档" if filename.lower().endswith(".md") else "文本文件"
    return _truncate(f"[{ext_label}: {filename}]\n{text}")
