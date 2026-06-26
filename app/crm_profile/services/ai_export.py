"""Generate editable AI reply exports for docx and pdf."""
from __future__ import annotations

import re
import textwrap
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from zoneinfo import ZoneInfo

import fitz
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ...config import settings

DOC_BODY_LINE_SPACING = 1.38
DOC_LIST_LINE_SPACING = 1.30
DOC_CALLOUT_LINE_SPACING = 1.36
PDF_BODY_LINE_GAP = 5.5
PDF_TIGHT_LINE_GAP = 4.5
PDF_CALLOUT_LINE_HEIGHT = 15.5


@dataclass
class ExportFile:
    content: bytes
    media_type: str
    filename: str


@dataclass
class MarkdownBlock:
    kind: str
    text: str
    level: int = 0


def build_ai_reply_export(
    *,
    content: str,
    export_format: str,
    title: str | None = None,
    customer_name: str | None = None,
) -> ExportFile:
    """Build a formatted export file from one AI assistant reply."""
    clean_content = (content or "").strip()
    if not clean_content:
        raise ValueError("导出内容不能为空")
    safe_title = (title or "AI教练回复").strip()[:80] or "AI教练回复"
    blocks = parse_markdown_blocks(clean_content)
    if export_format == "docx":
        data = _build_docx(blocks, safe_title, customer_name)
        return ExportFile(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{_safe_filename(safe_title)}.docx",
        )
    if export_format == "pdf":
        data = _build_pdf(blocks, safe_title, customer_name)
        return ExportFile(
            content=data,
            media_type="application/pdf",
            filename=f"{_safe_filename(safe_title)}.pdf",
        )
    raise ValueError("不支持的导出格式")


def parse_markdown_blocks(content: str) -> list[MarkdownBlock]:
    """Parse the subset of markdown we need for readable exports."""
    blocks: list[MarkdownBlock] = []
    lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    paragraph: list[str] = []
    code_lines: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
        if paragraph:
            text = " ".join(line.strip() for line in paragraph if line.strip())
            if text:
                blocks.append(MarkdownBlock("paragraph", _strip_inline_markdown(text)))
            paragraph.clear()

    def flush_code() -> None:
        if code_lines:
            blocks.append(MarkdownBlock("code", "\n".join(code_lines).strip("\n")))
            code_lines.clear()

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            blocks.append(MarkdownBlock("heading", _strip_inline_markdown(heading.group(2)), len(heading.group(1))))
            continue
        if re.match(r"^(\*{3,}|-{3,}|_{3,})$", stripped):
            flush_paragraph()
            blocks.append(MarkdownBlock("separator", ""))
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        ordered = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if ordered:
            flush_paragraph()
            blocks.append(MarkdownBlock("ordered", _strip_inline_markdown(ordered.group(2)), int(ordered.group(1))))
            continue
        if bullet:
            flush_paragraph()
            blocks.append(MarkdownBlock("bullet", _strip_inline_markdown(bullet.group(1))))
            continue
        quote = re.match(r"^>\s?(.+)$", stripped)
        if quote:
            flush_paragraph()
            blocks.append(MarkdownBlock("quote", _strip_inline_markdown(quote.group(1))))
            continue
        paragraph.append(stripped)
    flush_paragraph()
    flush_code()
    return blocks or [MarkdownBlock("paragraph", _strip_inline_markdown(content))]


def _build_docx(blocks: list[MarkdownBlock], title: str, customer_name: str | None) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    _set_doc_default_font(doc)
    title_para = doc.add_paragraph()
    title_para.style = doc.styles["Title"]
    title_run = title_para.add_run(title)
    _set_run_font(title_run, 18, bold=True, color=RGBColor(31, 41, 55))

    meta = doc.add_paragraph()
    meta_run = meta.add_run(_export_meta(customer_name))
    _set_run_font(meta_run, 9, color=RGBColor(107, 114, 128))

    for block in blocks:
        if block.kind == "heading":
            para = doc.add_paragraph()
            run = para.add_run(block.text)
            size = 15 if block.level <= 1 else 13
            _set_run_font(run, size, bold=True, color=RGBColor(17, 24, 39))
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(5)
            para.paragraph_format.line_spacing = 1.18
        elif block.kind == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(block.text)
            _set_run_font(run, 10.5)
            para.paragraph_format.line_spacing = DOC_LIST_LINE_SPACING
            para.paragraph_format.space_after = Pt(3)
        elif block.kind == "ordered":
            para = doc.add_paragraph(style="List Number")
            run = para.add_run(block.text)
            _set_run_font(run, 10.5)
            para.paragraph_format.line_spacing = DOC_LIST_LINE_SPACING
            para.paragraph_format.space_after = Pt(3)
        elif block.kind == "quote":
            label = doc.add_paragraph()
            label.paragraph_format.space_before = Pt(8)
            label.paragraph_format.space_after = Pt(0)
            _shade_paragraph(label, "EFF6FF")
            label_run = label.add_run("重点提示")
            _set_run_font(label_run, 9, bold=True, color=RGBColor(37, 99, 235))
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.08)
            para.paragraph_format.line_spacing = DOC_CALLOUT_LINE_SPACING
            para.paragraph_format.space_after = Pt(8)
            run = para.add_run(block.text)
            _set_run_font(run, 10, color=RGBColor(75, 85, 99))
            _shade_paragraph(para, "F8FBFF")
        elif block.kind == "code":
            label = doc.add_paragraph()
            label.paragraph_format.space_before = Pt(8)
            label.paragraph_format.space_after = Pt(0)
            _shade_paragraph(label, "EEF2FF")
            label_run = label.add_run("客户话术")
            _set_run_font(label_run, 9, bold=True, color=RGBColor(79, 70, 229))
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.08)
            para.paragraph_format.line_spacing = DOC_CALLOUT_LINE_SPACING
            para.paragraph_format.space_after = Pt(8)
            _shade_paragraph(para, "F8FAFC")
            for idx, line in enumerate(block.text.splitlines() or [""]):
                if idx:
                    para.add_run().add_break()
                run = para.add_run(line)
                _set_run_font(run, 10, color=RGBColor(31, 41, 55))
        elif block.kind == "separator":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            _add_bottom_border(para, "E5E7EB")
        else:
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = DOC_BODY_LINE_SPACING
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(block.text)
            _set_run_font(run, 10.5, color=RGBColor(55, 65, 81))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_pdf(blocks: list[MarkdownBlock], title: str, customer_name: str | None) -> bytes:
    doc = fitz.open()
    fontfile = _find_cjk_font()
    fontname = "msyh" if fontfile else "helv"
    page = doc.new_page(width=595, height=842)
    margin_x = 48
    y = 54
    bottom = 790

    def add_page() -> None:
        nonlocal page, y
        _draw_footer(page, len(doc))
        page = doc.new_page(width=595, height=842)
        y = 54

    def draw_line(text: str, *, size: float = 11, color: tuple[float, float, float] = (0.22, 0.25, 0.31), x_offset: float = 0, line_gap: float = PDF_BODY_LINE_GAP) -> None:
        nonlocal y
        if y > bottom:
            add_page()
        page.insert_text(
            (margin_x + x_offset, y),
            text,
            fontsize=size,
            fontname=fontname,
            fontfile=fontfile,
            color=color,
        )
        y += size + line_gap

    def draw_separator() -> None:
        nonlocal y
        if y > bottom - 12:
            add_page()
        page.draw_line((margin_x, y), (548, y), color=(0.86, 0.88, 0.91), width=0.6)
        y += 14

    def draw_callout(label: str, lines: list[str], *, accent: tuple[float, float, float], header_fill: tuple[float, float, float], body_fill: tuple[float, float, float]) -> None:
        nonlocal y
        pending = lines or [""]
        line_height = PDF_CALLOUT_LINE_HEIGHT
        while pending:
            available = int((bottom - y - 48) / line_height)
            if available <= 0:
                add_page()
                continue
            chunk = pending[:available]
            pending = pending[available:]
            header_h = 26
            body_h = len(chunk) * line_height + 18
            header_rect = fitz.Rect(margin_x, y, 548, y + header_h)
            body_rect = fitz.Rect(margin_x, y + header_h, 548, y + header_h + body_h)
            page.draw_rect(header_rect, color=header_fill, fill=header_fill, width=0)
            page.draw_rect(body_rect, color=body_fill, fill=body_fill, width=0)
            page.draw_rect(fitz.Rect(margin_x, y, margin_x + 3, y + header_h + body_h), color=accent, fill=accent, width=0)
            page.insert_text(
                (margin_x + 12, y + 18),
                label,
                fontsize=9,
                fontname=fontname,
                fontfile=fontfile,
                color=accent,
            )
            text_y = y + header_h + 16
            for line in chunk:
                page.insert_text(
                    (margin_x + 12, text_y),
                    line,
                    fontsize=10,
                    fontname=fontname,
                    fontfile=fontfile,
                    color=(0.18, 0.22, 0.28),
                )
                text_y += line_height
            y += header_h + body_h + 10

    draw_line(title, size=18, color=(0.08, 0.10, 0.14), line_gap=12)
    draw_line(_export_meta(customer_name), size=9, color=(0.42, 0.45, 0.50), line_gap=18)

    for block in blocks:
        if block.kind == "heading":
            y += 5
            for line in _wrap_text(block.text, 28 if block.level <= 1 else 34):
                draw_line(line, size=14 if block.level <= 1 else 12.5, color=(0.07, 0.09, 0.13), line_gap=7)
        elif block.kind == "bullet":
            for idx, line in enumerate(_wrap_text(block.text, 42)):
                draw_line(("• " if idx == 0 else "  ") + line, size=10.5, x_offset=8, line_gap=PDF_TIGHT_LINE_GAP)
        elif block.kind == "ordered":
            for idx, line in enumerate(_wrap_text(block.text, 40)):
                prefix = f"{block.level}. " if idx == 0 else "   "
                draw_line(prefix + line, size=10.5, x_offset=4, line_gap=PDF_TIGHT_LINE_GAP)
        elif block.kind == "quote":
            draw_callout(
                "重点提示",
                _wrap_text(block.text, 42),
                accent=(0.15, 0.39, 0.92),
                header_fill=(0.93, 0.96, 1.00),
                body_fill=(0.98, 0.99, 1.00),
            )
        elif block.kind == "code":
            code_lines: list[str] = []
            for line in block.text.splitlines() or [""]:
                code_lines.extend(_wrap_text(line, 42))
            draw_callout(
                "客户话术",
                code_lines,
                accent=(0.31, 0.27, 0.90),
                header_fill=(0.94, 0.95, 1.00),
                body_fill=(0.99, 0.99, 1.00),
            )
        elif block.kind == "separator":
            draw_separator()
        else:
            for line in _wrap_text(block.text, 42):
                draw_line(line, size=10.5, line_gap=PDF_BODY_LINE_GAP)
            y += 4

    _draw_footer(page, len(doc))
    data = doc.tobytes(garbage=4, deflate=True)
    doc.close()
    return data


def _set_doc_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_run_font(run, size: float, *, bold: bool = False, color: RGBColor | None = None, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _shade_paragraph(para, fill: str) -> None:
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _add_bottom_border(para, color: str) -> None:
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^[#>\s]+", "", text)
    return text.strip()


def _wrap_text(text: str, width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    for raw in text.splitlines():
        wrapped = textwrap.wrap(raw, width=width, break_long_words=True, replace_whitespace=False)
        lines.extend(wrapped or [""])
    return lines


def _draw_footer(page: fitz.Page, page_number: int) -> None:
    page.insert_text((280, 818), str(page_number), fontsize=8, color=(0.55, 0.58, 0.64))


def _find_cjk_font() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


def _export_meta(customer_name: str | None) -> str:
    now = datetime.now(ZoneInfo(settings.default_timezone)).strftime("%Y-%m-%d %H:%M")
    name = (customer_name or "当前客户").strip() or "当前客户"
    return f"{name} · 导出时间 {now}"


def _safe_filename(value: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]+', "_", value).strip(" ._")
    return name[:64] or "AI教练回复"
